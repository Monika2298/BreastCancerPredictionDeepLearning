from flask import Flask, render_template, request,send_file
from werkzeug.utils import secure_filename
import amy
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import io
import tensorflow as tf
from keras.models import model_from_json
from keras.models import load_model
from docx import Document
import os
from plotly.offline import download_plotlyjs,init_notebook_mode,iplot
import cufflinks as cf
init_notebook_mode(connected=True)
cf.go_offline()
app = Flask(__name__,template_folder='/Users/manish/Desktop/BE project')





@app.route('/upload')
def upload_file():
    return render_template('upload.html')

@app.route('/plots',methods=['GET','POST'])
def plots():
    return render_template('plots.html')


@app.route('/uploader', methods = ['GET', 'POST'])
def uploader():
    
    if request.method == 'POST':
        f = request.files['file']
        #f.save(secure_filename(f.filename))
        df=pd.read_csv(f)
    return render_template("submit.html", content=df)
   
@app.route('/result', methods = ['GET', 'POST'])
def result():
        
    f= request.files['file']
    json_file = open('model22.json', 'r')
    loaded_model_json = json_file.read()
    json_file.close()
    loaded_model = model_from_json(loaded_model_json)
        # load weights into new model
    loaded_model.load_weights("model22.h5")

    loaded_model.compile(loss='binary_crossentropy', optimizer='rmsprop', metrics=['accuracy'])

    loaded_model.save("model22.h5")
    new_model=load_model('model22.h5')

    
    from sklearn.preprocessing import StandardScaler
    scalar=StandardScaler()

    col_list=['radius_mean', 'texture_mean', 'perimeter_mean', 'area_mean',
       'smoothness_mean', 'compactness_mean', 'concavity_mean',
       'concave points_mean', 'symmetry_mean', 'fractal_dimension_mean',
       'radius_se', 'texture_se', 'perimeter_se', 'area_se', 'smoothness_se',
       'compactness_se', 'concavity_se', 'concave points_se', 'symmetry_se',
       'fractal_dimension_se', 'radius_worst', 'texture_worst',
       'perimeter_worst', 'area_worst', 'smoothness_worst',
       'compactness_worst', 'concavity_worst', 'concave points_worst',
       'symmetry_worst', 'fractal_dimension_worst']
    df = pd.read_csv(f,usecols=col_list)
    x=scalar.fit_transform(df)
    output=new_model.predict_classes(x)
    output1=new_model.predict(x)

    df1 = pd.read_csv("inputfile.csv")

    df1['prediction'] = output
    df1['Probability']= output1

    dictionary={ 1:'Benign', 0:'Malignant'}
    df1.prediction=[dictionary[item] for item in df1.prediction]
    df1.to_csv("result3.csv", index=False)

    table= pd.read_csv("result3.csv") 
    p=table[['Patient ID','prediction','Probability']]
    #s=df1['prediction'].value_counts().iplot(kind='bar',title='Number of Malignant and Benign Patients(Countplot)')


  
    return render_template("result.html", content=p)
    #return s

    

@app.route('/download')
def download_file():
	path = "result3.csv"
	return send_file(path, as_attachment=True,cache_timeout=0)




@app.route('/patient', methods = ['GET', 'POST'])
def patient():
    data = pd.read_csv("result3.csv", index_col ="Patient ID")
    if request.method=='POST':
        y=request.form['ID']
        global n
        global doc
        n=int(y)
        if(n in data.index.values):
            doc=data.loc[n]


            
            document=Document()
            h=document.add_heading('Patient ID-',0)
            h.add_run(str(n),0)

            p=document.add_paragraph(str(doc))


            for paragraph in document.paragraphs:
                if 'Malignant' in paragraph.text:
                    x=document.add_paragraph('The diagnosis for the patient is ')
                    x.add_run(' MALIGNANT').bold = True
                    print(x)
                    m="Malignant"
        
        
        
                if 'Benign' in paragraph.text:
                    x=document.add_paragraph('The diagnosis for the patient is ')
                    x.add_run(' BENIGN').bold = True
                    print(x)
                    m="Benign"

            document.save(str(n)+".docx")
            return render_template("patient.html", content=doc,id=y,d=m)

        else:
            s="Not Valid"
            return render_template("invalid.html",id=s)

#report
    
    

@app.route('/downloadreport', methods = ['GET', 'POST'])
def download_patient():
    global n
    return send_file(str(n)+".docx", as_attachment=True, cache_timeout=0)
    


       
    

        
if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)

